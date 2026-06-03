import re
import shutil
import subprocess
import os
from bs4 import BeautifulSoup
from bs4.element import NavigableString, Tag
import markdown
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import latex2mathml.converter
import mathml2omml

class DocumentConverter:
    def __init__(self):
        self.equations = {}
        self.eq_counter = 0

    def _replace_math(self, match, is_block):
        self.eq_counter += 1
        placeholder = f"EQ_PLACEHOLDER_{self.eq_counter}"
        self.equations[placeholder] = {
            "latex": match.group(1).strip(),
            "is_block": is_block
        }
        return placeholder

    # ------------------------------------------------------------------
    # AI-SPECIFIC PREPROCESSORS
    # ------------------------------------------------------------------

    def _preprocess_deepseek(self, text):
        """
        DeepSeek uses \( \) for inline and \[ \] for display math.
        Convert to $ $ and $$ $$ so the standard pipeline can handle them.
        Also normalises fraction-style text that DeepSeek sometimes produces.
        """
        # Display math \[ ... \] → $$ ... $$
        text = re.sub(r'\\\[([^\]]+?)\\\]', r'$$\1$$', text, flags=re.DOTALL)
        # Inline math \( ... \) → $ ... $
        text = re.sub(r'\\\(([^\)]+?)\\\)', r'$\1$', text, flags=re.DOTALL)
        return text

    def _preprocess_google_ia(self, text):
        """
        Google Modo IA quirks:
        - Uses 30-dash horizontal rules: ------------------------------ → ---
        - Has numbered footnote references [N] inline → remove trailing ones
        - Has URL-style reference list at bottom → strip it
        """
        # Normalise horizontal rules
        text = re.sub(r'-{10,}', '---', text)
        # Remove inline reference markers like [1] [2] at end of sentences
        text = re.sub(r'\s*\[(\d+)\]', '', text)
        # Remove trailing URL reference block (lines starting with [N] https://...)
        text = re.sub(r'\n\[\d+\]\s+https?://\S+[^\n]*', '', text)
        return text

    def _preprocess_notebooklm(self, text):
        """
        NotebookLM Chat quirks:
        - Has a conversational tail offering Deep Research or referencing no sources.
          Strip the trailing disclaimer paragraphs.
        - Tables use ':---' alignment markers (fully valid Markdown, no special fix needed).
        """
        # Remove "NotebookLM funciona mejor cuando..." disclaimer block
        disclaimer_patterns = [
            r'\nNotebookLM funciona mejor cuando.*',
            r'\nNotebook.*Deep Research.*',
            r'\n¿Te gustaría que realice una.*',
        ]
        for p in disclaimer_patterns:
            text = re.sub(p, '', text, flags=re.DOTALL | re.IGNORECASE)
        return text

    def _preprocess_generic(self, text):
        """
        Common fixes for all IAs:
        - Ensure $$ equations are on their own line (some IAs put them inline mid-sentence)
        - Normalise Windows/Mac line-endings
        """
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        # If $$ ... $$ is NOT already on its own line, inject newlines around it
        text = re.sub(r'([^\n])\$\$', r'\1\n$$', text)
        text = re.sub(r'\$\$([^\n])', r'$$\n\1', text)
        return text

    def preprocess_for_ai(self, text, ai_source="Gemini AI"):
        """
        Entry point: applies the generic fix first, then the AI-specific one.
        """
        text = self._preprocess_generic(text)
        source = ai_source.lower()
        if "deepseek" in source:
            text = self._preprocess_deepseek(text)
        elif "google" in source:
            text = self._preprocess_google_ia(text)
        elif "notebooklm" in source:
            text = self._preprocess_notebooklm(text)
        # Gemini AI and ChatGPT produce fairly standard Markdown – no extra fixes needed
        return text

    # ------------------------------------------------------------------
    # MATH PROTECTION
    # ------------------------------------------------------------------

    def preprocess_markdown(self, text):
        """
        Extracts equations and replaces them with placeholders so the Markdown
        parser does not mangle dollar signs, asterisks, or underscores inside them.
        """
        # Block equations first (longer patterns take priority)
        text = re.sub(r'\$\$([\s\S]+?)\$\$', lambda m: self._replace_math(m, True) + '\n', text)
        text = re.sub(r'\\\[([\s\S]+?)\\\]',  lambda m: self._replace_math(m, True) + '\n', text)
        # Inline equations
        text = re.sub(r'\$([^\$\n]+?)\$', lambda m: self._replace_math(m, False), text)
        text = re.sub(r'\\\(([\s\S]+?)\\\)', lambda m: self._replace_math(m, False), text)
        return text

    # ------------------------------------------------------------------
    # MATH RENDERING
    # ------------------------------------------------------------------

    def generate_omml(self, latex_string):
        try:
            mathml = latex2mathml.converter.convert(latex_string)
            mathml = mathml.replace('display="block"', '')
            omml_str = mathml2omml.convert(mathml)
            if "<m:oMath>" in omml_str:
                omml_str = omml_str.replace("<m:oMath>",
                    '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">', 1)
            elif "<m:oMathPara>" in omml_str:
                omml_str = omml_str.replace("<m:oMathPara>",
                    '<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">', 1)
            return etree.fromstring(omml_str.encode('utf-8'))
        except Exception as e:
            print(f"Error converting equation: {latex_string}\n{e}")
            r = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
            t = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
            t.text = f"[Equ Failed: {latex_string}]"
            r.append(t)
            return r

    # ------------------------------------------------------------------
    # HTML NODE PROCESSOR (inline elements + equations)
    # ------------------------------------------------------------------

    def _process_node(self, node, paragraph):
        """
        Walk an HTML node tree and add runs/equations to `paragraph`.
        Rule: paragraph alignment is set to CENTER only when a BLOCK equation
        is encountered. Regular text paragraphs keep LEFT (default) alignment.
        """
        if isinstance(node, NavigableString):
            text = str(node)
            if "EQ_PLACEHOLDER_" in text:
                parts = re.split(r'(EQ_PLACEHOLDER_\d+)', text)
                for part in parts:
                    if part.startswith("EQ_PLACEHOLDER_") and part in self.equations:
                        eq_data = self.equations[part]
                        omml_elem = self.generate_omml(eq_data["latex"])
                        if eq_data["is_block"]:
                            # Only block equations centre the paragraph
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph._element.append(omml_elem)
                    elif part:
                        paragraph.add_run(part)
            else:
                if text.strip() or text == " ":
                    paragraph.add_run(text)
            return

        if isinstance(node, Tag):
            pass  # handled via children below

        for child in node:
            if isinstance(child, NavigableString):
                runtext = str(child)
                parts = re.split(r'(EQ_PLACEHOLDER_\d+)', runtext)
                for part in parts:
                    if part.startswith("EQ_PLACEHOLDER_") and part in self.equations:
                        eq_data = self.equations[part]
                        omml_elem = self.generate_omml(eq_data["latex"])
                        if eq_data["is_block"]:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph._element.append(omml_elem)
                    elif part:
                        run = paragraph.add_run(part)
                        if node.name in ('strong', 'b'):
                            run.bold = True
                        if node.name in ('em', 'i'):
                            run.italic = True
                        if node.name == 'code':
                            run.font.name = 'Courier New'
            else:
                self._process_node(child, paragraph)

    # ------------------------------------------------------------------
    # TABLE PROCESSOR
    # ------------------------------------------------------------------

    def _set_cell_shading(self, cell, fill_color):
        """Apply background colour to a table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_color)
        tcPr.append(shd)

    def _process_table(self, table_tag, doc, font_name, font_size):
        """
        Converts an HTML <table> element from BeautifulSoup into a Word table.
        Handles headers (<th>) and body rows (<td>).
        Equations inside cells are rendered correctly.
        """
        rows = table_tag.find_all('tr')
        if not rows:
            return

        # Count columns from the first row
        first_row_cells = rows[0].find_all(['th', 'td'])
        num_cols = len(first_row_cells)
        if num_cols == 0:
            return

        tbl = doc.add_table(rows=len(rows), cols=num_cols)
        tbl.style = 'Table Grid'

        for row_idx, tr in enumerate(rows):
            cells = tr.find_all(['th', 'td'])
            is_header = all(c.name == 'th' for c in cells)

            for col_idx, cell_tag in enumerate(cells):
                if col_idx >= num_cols:
                    break
                cell = tbl.cell(row_idx, col_idx)
                # Clear default empty paragraph
                cell.paragraphs[0].clear()
                para = cell.paragraphs[0]

                # Style: header row
                if is_header or cell_tag.name == 'th':
                    self._set_cell_shading(cell, '4472C4')  # blue header
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_style = {'bold': True, 'color': RGBColor(0xFF, 0xFF, 0xFF)}
                else:
                    run_style = {}

                # Process content (text + equations)
                for child in cell_tag.children:
                    if isinstance(child, NavigableString):
                        txt = str(child)
                        parts = re.split(r'(EQ_PLACEHOLDER_\d+)', txt)
                        for part in parts:
                            if part.startswith("EQ_PLACEHOLDER_") and part in self.equations:
                                eq_data = self.equations[part]
                                omml_elem = self.generate_omml(eq_data["latex"])
                                para._element.append(omml_elem)
                            elif part:
                                run = para.add_run(part)
                                run.font.name = font_name
                                run.font.size = Pt(font_size)
                                if run_style.get('bold'):
                                    run.bold = True
                                if 'color' in run_style:
                                    run.font.color.rgb = run_style['color']
                    else:
                        # Handle nested tags (strong, em, code, etc.)
                        self._process_node(child, para)
                        # Re-apply header styles to any runs added
                        if is_header or cell_tag.name == 'th':
                            for run in para.runs:
                                run.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                run.font.name = font_name
                                run.font.size = Pt(font_size)

        # Set column widths to auto-fit content
        for col in tbl.columns:
            for cell in col.cells:
                cell.width = Inches(6.0 / num_cols)

        doc.add_paragraph()  # spacing after table

    # ------------------------------------------------------------------
    # MAIN CONVERSION PIPELINE
    # ------------------------------------------------------------------

    def convert_markdown_to_docx(self, md_text, output_path, font_name="Arial", font_size=11, ai_source="Gemini AI"):
        # Reset state
        self.equations = {}
        self.eq_counter = 0

        # 1. AI-specific preprocessing
        md_text = self.preprocess_for_ai(md_text, ai_source)

        # 2. Protect Math (extract equations → placeholders)
        protected_md = self.preprocess_markdown(md_text)

        # 3. Markdown to HTML (with tables extension)
        html = markdown.markdown(protected_md, extensions=['tables', 'fenced_code', 'sane_lists'])

        # 4. Parse HTML
        soup = BeautifulSoup(html, 'html.parser')

        # 5. Build DOCX
        doc = Document()

        # Apply chosen font to all base styles
        normal_style = doc.styles['Normal']
        normal_style.font.name = font_name
        normal_style.font.size = Pt(font_size)

        for s in doc.styles:
            if hasattr(s, 'font') and s.font is not None:
                s.font.name = font_name
                if s.element.rPr is not None and s.element.rPr.rFonts is not None:
                    for theme_attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:cstheme']:
                        if s.element.rPr.rFonts.get(qn(theme_attr)):
                            s.element.rPr.rFonts.set(qn(theme_attr), '')

        for element in soup.contents:
            if not isinstance(element, Tag):
                continue

            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                p = doc.add_heading('', level=level)
                self._process_node(element, p)

            elif element.name == 'p':
                # Check if this paragraph contains ONLY a block equation
                # In that case it should be centred; otherwise LEFT-aligned.
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # default — may be overridden by block eq
                self._process_node(element, p)

            elif element.name == 'blockquote':
                p = doc.add_paragraph(style='Quote')
                self._process_node(element, p)

            elif element.name in ('ul', 'ol'):
                list_style = 'List Bullet' if element.name == 'ul' else 'List Number'
                for li in element.find_all('li', recursive=False):
                    p = doc.add_paragraph(style=list_style)
                    self._process_node(li, p)

            elif element.name == 'table':
                self._process_table(element, doc, font_name, font_size)

            elif element.name == 'hr':
                # Horizontal rule → thin paragraph border
                p = doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), 'AAAAAA')
                pBdr.append(bottom)
                pPr.append(pBdr)

            else:
                p = doc.add_paragraph()
                self._process_node(element, p)

        doc.save(output_path)
        return output_path

    def convert_with_pandoc(self, md_text, output_path):
        temp_md = "temp_input.md"
        with open(temp_md, "w", encoding="utf-8") as f:
            f.write(md_text)
        try:
            subprocess.run(["pandoc", "-f", "markdown", "-t", "docx", "-o", output_path, temp_md], check=True)
            os.remove(temp_md)
            return True, output_path
        except Exception as e:
            if os.path.exists(temp_md):
                os.remove(temp_md)
            return False, str(e)
